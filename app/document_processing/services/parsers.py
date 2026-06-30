import io
import re
import logging
from abc import ABC, abstractmethod
from typing import Dict, Any, Type
from pydantic import BaseModel

logger = logging.getLogger(__name__)


class ExtractedText(BaseModel):
    """Unified output structure for all document parsers."""
    raw_text: str
    page_count: int
    character_count: int
    word_count: int
    parser_used: str
    metadata: Dict[str, Any]


class BaseParser(ABC):
    """Abstract Base Class for all document parsers."""

    @abstractmethod
    def parse(self, file_bytes: bytes, filename: str) -> ExtractedText:
        """Parse raw file bytes into standard ExtractedText structure."""
        pass


class PDFParser(BaseParser):
    """Parser for PDF files using pypdf."""

    def parse(self, file_bytes: bytes, filename: str) -> ExtractedText:
        logger.info(f"PDFParser started parsing {filename}")
        try:
            import pypdf
        except ImportError:
            raise RuntimeError("pypdf is required to parse PDF documents.")

        pdf_file = io.BytesIO(file_bytes)
        try:
            reader = pypdf.PdfReader(pdf_file)
        except Exception as e:
            logger.error(f"Corrupted or invalid PDF {filename}: {e}")
            raise ValueError(f"Could not parse corrupted PDF file: {e}")

        if reader.is_encrypted:
            logger.warning(f"Encrypted PDF detected: {filename}. Attempting decryption with empty password.")
            try:
                reader.decrypt("")
            except Exception as e:
                raise ValueError(f"PDF file is encrypted and cannot be decrypted: {e}")

        pages_text = []
        page_count = len(reader.pages)
        
        for idx, page in enumerate(reader.pages):
            try:
                text = page.extract_text() or ""
                pages_text.append(text)
                logger.debug(f"Successfully processed page {idx+1}/{page_count} for {filename}")
            except Exception as e:
                logger.warning(f"Error extracting text from page {idx+1} of {filename}: {e}")
                pages_text.append("")

        full_text = "\n".join(pages_text)
        
        return ExtractedText(
            raw_text=full_text,
            page_count=page_count,
            character_count=len(full_text),
            word_count=len(full_text.split()),
            parser_used="pypdf",
            metadata={"encrypted": reader.is_encrypted}
        )


class DOCXParser(BaseParser):
    """Parser for Word DOCX files using python-docx."""

    def parse(self, file_bytes: bytes, filename: str) -> ExtractedText:
        logger.info(f"DOCXParser started parsing {filename}")
        try:
            import docx
        except ImportError:
            raise RuntimeError("python-docx is required to parse DOCX documents.")

        docx_file = io.BytesIO(file_bytes)
        try:
            doc = docx.Document(docx_file)
        except Exception as e:
            logger.error(f"Corrupted or invalid DOCX {filename}: {e}")
            raise ValueError(f"Could not parse corrupted DOCX file: {e}")

        elements = []
        
        # 1. Process paragraphs (includes headings and list items)
        for para in doc.paragraphs:
            if para.text.strip():
                elements.append(para.text)

        # 2. Process tables (text-only)
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    elements.append(" | ".join(row_text))

        full_text = "\n".join(elements)
        
        return ExtractedText(
            raw_text=full_text,
            page_count=1,  # DOCX lacks physical pages without rendering engine
            character_count=len(full_text),
            word_count=len(full_text.split()),
            parser_used="python-docx",
            metadata={"paragraphs_count": len(doc.paragraphs), "tables_count": len(doc.tables)}
        )


class TXTParser(BaseParser):
    """Parser for plaintext files."""

    def parse(self, file_bytes: bytes, filename: str) -> ExtractedText:
        logger.info(f"TXTParser started parsing {filename}")
        
        # Attempt decoding with UTF-8, fallback to latin-1
        try:
            full_text = file_bytes.decode("utf-8")
        except UnicodeDecodeError:
            logger.warning(f"UTF-8 decode failed for {filename}, falling back to ISO-8859-1")
            full_text = file_bytes.decode("iso-8859-1")

        return ExtractedText(
            raw_text=full_text,
            page_count=1,
            character_count=len(full_text),
            word_count=len(full_text.split()),
            parser_used="native_txt",
            metadata={}
        )


class MarkdownParser(BaseParser):
    """Parser for Markdown files, cleaning basic syntax while preserving structure."""

    def parse(self, file_bytes: bytes, filename: str) -> ExtractedText:
        logger.info(f"MarkdownParser started parsing {filename}")
        
        # Safe string decode
        try:
            raw_markdown = file_bytes.decode("utf-8")
        except UnicodeDecodeError:
            raw_markdown = file_bytes.decode("iso-8859-1")

        # Clean Markdown formatting elements (headings, bold, italic, links, blockquotes)
        text = raw_markdown
        
        # 1. Links: [label](url) -> label
        text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)
        
        # 2. Headings: # Heading -> Heading
        text = re.sub(r'^#+\s+(.+)$', r'\1', text, flags=re.MULTILINE)
        
        # 3. Bold/Italics: **text** or *text* -> text
        text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
        text = re.sub(r'\*([^*]+)\*', r'\1', text)
        text = re.sub(r'__([^_]+)__', r'\1', text)
        text = re.sub(r'_([^_]+)_', r'\1', text)

        # 4. Code Blocks: ```code``` -> code
        text = re.sub(r'```[\s\S]*?```', lambda m: m.group(0).replace('```', ''), text)
        
        # 5. Inline Code: `code` -> code
        text = re.sub(r'`([^`]+)`', r'\1', text)

        return ExtractedText(
            raw_text=text,
            page_count=1,
            character_count=len(text),
            word_count=len(text.split()),
            parser_used="native_markdown",
            metadata={}
        )


class ParserFactory:
    """Factory class to select and initialize the correct document parser based on file extension."""

    _parsers: Dict[str, Type[BaseParser]] = {
        ".pdf": PDFParser,
        ".docx": DOCXParser,
        ".txt": TXTParser,
        ".md": MarkdownParser,
    }

    @classmethod
    def get_parser(cls, filename: str) -> BaseParser:
        import os
        _, ext = os.path.splitext(filename.lower())
        
        parser_cls = cls._parsers.get(ext)
        if not parser_cls:
            raise ValueError(f"No parser available for extension: {ext}")
            
        return parser_cls()
